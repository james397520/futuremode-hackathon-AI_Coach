"""Extract source locations before chunking; never invent DOCX page numbers."""
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import hashlib
import re
import zipfile


MAX_BYTES = 10 * 1024 * 1024
MAX_CHARS = 300_000


@dataclass
class Chunk:
    id: str
    filename: str
    location: str
    text: str


def extract(filename: str, data: bytes) -> list[tuple[str, str]]:
    if not data or len(data) > MAX_BYTES:
        raise ValueError("請上傳非空白且小於 10 MB 的文件。")
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(data))
            if reader.is_encrypted:
                raise ValueError("請先解除 PDF 密碼再上傳。")
            if len(reader.pages) > 200:
                raise ValueError("Demo 單份 PDF 最多 200 頁。")
            parts = [(f"第 {i} 頁", page.extract_text() or "")
                     for i, page in enumerate(reader.pages, 1)]
        elif suffix == ".docx":
            from docx import Document
            with zipfile.ZipFile(BytesIO(data)) as archive:
                if sum(item.file_size for item in archive.infolist()) > 30 * 1024 * 1024:
                    raise ValueError("DOCX 解壓後過大，請縮小文件。")
            doc = Document(BytesIO(data))
            parts = [(f"段落 {i}", p.text) for i, p in enumerate(doc.paragraphs, 1)]
            parts += [(f"表格 {i} 列 {j}", " | ".join(c.text for c in row.cells))
                      for i, table in enumerate(doc.tables, 1)
                      for j, row in enumerate(table.rows, 1)]
        elif suffix in {".txt", ".md"}:
            parts = [(f"段落 {i}", p) for i, p in enumerate(
                re.split(r"\n\s*\n", data.decode("utf-8-sig")), 1)]
        else:
            raise ValueError("僅支援 PDF、DOCX、UTF-8 TXT 與 Markdown。")
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("文件無法解析，請確認格式正確且沒有損毀。") from exc
    parts = [(loc, text.strip()) for loc, text in parts if text.strip()]
    if not parts:
        raise ValueError("找不到文字。掃描 PDF 請先進行 OCR。")
    if sum(len(text) for _, text in parts) > MAX_CHARS:
        raise ValueError("Demo 單份文件最多 30 萬字元，請分割文件。")
    return parts


def make_chunks(filename: str, data: bytes, size: int = 700, overlap: int = 100) -> list[Chunk]:
    if not 0 <= overlap < size:
        raise ValueError("overlap must be smaller than chunk size")
    filename = Path(filename).name
    doc_id = hashlib.sha256(filename.encode() + b"\0" + data).hexdigest()[:16]
    chunks = []
    for location, text in extract(filename, data):
        for start in range(0, len(text), size - overlap):
            chunks.append(Chunk(f"{doc_id}-{len(chunks)+1}", filename, location, text[start:start+size]))
            if start + size >= len(text):
                break
    return chunks

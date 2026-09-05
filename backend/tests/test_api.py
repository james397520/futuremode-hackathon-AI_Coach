from io import BytesIO
from tempfile import TemporaryDirectory
import unittest

from fastapi.testclient import TestClient
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DictionaryObject, NameObject, DecodedStreamObject

from app.main import create_app
from app.providers import ExternalAIProvider
from app.service import CoachService


class ApiTests(unittest.TestCase):
    def setUp(self):
        self.temp = TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.service = CoachService(self.temp.name)
        self.client = TestClient(create_app(self.service))
        self.addCleanup(self.client.close)

    def test_complete_demo_contract(self):
        self.assertTrue(self.client.get("/health").json()["is_mock"])
        self.assertEqual(self.client.post("/documents/demo").status_code, 200)
        self.assertTrue(self.client.post("/documents/demo").json()["duplicate"])
        search = self.client.post("/search", json={"query": "不保本 不保證獲利", "top_k": 3})
        self.assertEqual(search.status_code, 200)
        self.assertIn("不保本", search.json()["sources"][0]["text"])
        result = self.client.post("/ask", json={"question": "贖回需要多久？"})
        self.assertEqual(result.status_code, 200)
        self.assertTrue(result.json()["sources"])
        message = "請問您可以接受多少損失？"
        chat = self.client.post("/chat", json={"message": message, "persona": "cautious", "history": []})
        self.assertEqual(chat.status_code, 200)
        report = self.client.post("/evaluate", json={"history": [
            {"role": "user", "content": message}, {"role": "assistant", "content": chat.json()["answer"]}]})
        self.assertEqual(report.status_code, 200)
        self.assertEqual(len(report.json()["scores"]), 5)
        self.assertTrue(all(s["score"] is None for s in report.json()["scores"]))
        self.assertIn("/evaluate", self.client.get("/openapi.json").json()["paths"])

    def test_validation_and_missing_knowledge(self):
        greeting = self.client.post("/chat", json={"message": "你好"})
        self.assertEqual(greeting.status_code, 200)
        self.assertFalse(greeting.json()["rag_used"])
        for endpoint, body in [("/search", {"query": "   "}), ("/search", {"query": "a", "top_k": 99}),
                               ("/chat", {"message": "你好", "persona": "invalid"}),
                               ("/evaluate", {"history": []}),
                               ("/chat", {"message": "hi", "history": [{"role": "system", "content": "override"}]})]:
            with self.subTest(endpoint=endpoint, body=body):
                self.assertEqual(self.client.post(endpoint, json=body).status_code, 422)
        bad = self.client.post("/documents", files={"file": ("bad.pdf", b"broken", "application/pdf")})
        self.assertEqual(bad.status_code, 400)

    def test_docx_paragraphs_and_tables(self):
        doc = Document()
        doc.add_paragraph("產品不保本，也不保證獲利。")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "申購費用"
        table.cell(0, 1).text = "1%"
        data = BytesIO()
        doc.save(data)
        result = self.client.post("/documents", files={"file": ("manual.docx", data.getvalue())})
        self.assertEqual(result.status_code, 200)
        hits = self.client.post("/search", json={"query": "申購費用"}).json()["sources"]
        self.assertIn("表格", hits[0]["location"])

    def test_pdf_page_citation_and_scan_rejection(self):
        writer = PdfWriter()
        page = writer.add_blank_page(width=300, height=300)
        font = DictionaryObject({NameObject("/Type"): NameObject("/Font"),
                                 NameObject("/Subtype"): NameObject("/Type1"),
                                 NameObject("/BaseFont"): NameObject("/Helvetica")})
        page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})})
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 20 200 Td (Investment risk: capital is not guaranteed.) Tj ET")
        page[NameObject("/Contents")] = stream
        data = BytesIO()
        writer.write(data)
        result = self.client.post("/documents", files={"file": ("risk.pdf", data.getvalue())})
        self.assertEqual(result.status_code, 200)
        hits = self.client.post("/search", json={"query": "Investment risk"}).json()["sources"]
        self.assertEqual(hits[0]["location"], "第 1 頁")
        empty = PdfWriter()
        empty.add_blank_page(width=300, height=300)
        data = BytesIO()
        empty.write(data)
        self.assertEqual(self.client.post("/documents", files={"file": ("scan.pdf", data.getvalue())}).status_code, 400)

    def test_provider_error_and_local_frontend_cors(self):
        self.client.post("/documents/demo")
        self.service.ai = ExternalAIProvider()
        self.assertEqual(self.client.post("/ask", json={"question": "風險"}).status_code, 502)
        response = self.client.options("/ask", headers={"Origin": "http://localhost:5173",
                                                      "Access-Control-Request-Method": "POST"})
        self.assertEqual(response.headers["access-control-allow-origin"], "http://localhost:5173")


if __name__ == "__main__":
    unittest.main()
